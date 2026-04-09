from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm


TEMPLATE_PATH = Path(r"E:\表格\xzf毕设\第 4 章 概要设计.docx")
IMAGES_DIR = Path(r"E:\DownLoad\Idea-project\gaoxiaokeyan-front\docs\prototype\taskflow_images")
OUTPUT_PATH = Path(r"E:\表格\xzf毕设\第 4 章 概要设计-任务流图-无外框版.docx")


MODULES = [
    {
        "heading": "4.2.1 登录模块设计",
        "paragraphs": [
            "登录模块是系统全部业务功能的统一入口，所有用户在进入系统后均需先通过账号与密码完成身份认证。系统在校验登录信息成功后，将根据用户的角色自动跳转至对应默认页面；若登录失败，则返回相应提示信息，便于用户重新进行身份校验。",
            "由于本系统存在系统管理员、教师用户、财务管理员及多级审批人等角色，因此登录模块不仅承担用户认证功能，同时也负责后续业务页面访问权限的分发。该模块任务流图如图4-1所示：",
        ],
        "image": "fig4-1-login.png",
        "caption": "图4-1 登录模块任务流图",
    },
    {
        "heading": "4.2.2 首页工作台模块设计",
        "paragraphs": [
            "首页工作台模块主要面向系统管理员与教师用户，用于在进入系统后快速查看项目总数、进行中项目、待审批数量及已完成项目等关键统计信息。该页面同时展示最近通知与待办事项，为用户提供统一的信息汇总入口。",
            "首页工作台模块以数据展示为主，其核心流程为进入页面、加载项目数据、统计关键指标并完成工作台信息呈现。该模块任务流图如图4-2所示：",
        ],
        "image": "fig4-2-dashboard.png",
        "caption": "图4-2 首页工作台模块任务流图",
    },
    {
        "heading": "4.2.3 项目管理模块设计",
        "paragraphs": [
            "项目管理模块是系统核心业务模块之一，主要面向系统管理员与教师用户开放。用户可在该模块中完成项目的条件查询、分页浏览、新增、编辑、删除、导出以及跳转详情等操作，从而实现项目全生命周期管理。",
            "项目管理模块中的主要业务围绕项目列表展开，用户在筛选查询后可根据实际业务需求选择不同操作路径。该模块任务流图如图4-3所示：",
        ],
        "image": "fig4-3-project-list.png",
        "caption": "图4-3 项目管理模块任务流图",
    },
    {
        "heading": "4.2.4 项目详情与报销申请设计",
        "paragraphs": [
            "项目详情与报销申请模块用于展示单个项目的基础信息、预算明细与报销记录，主要面向系统管理员与教师用户开放。用户可在该页面中查看项目预算执行情况，并对当前项目提交报销申请、上传报销凭证与查看报销详情。",
            "该模块将项目查看与报销申请流程进行整合，可在同一页面中完成信息核对和业务提交。该模块任务流图如图4-4所示：",
        ],
        "image": "fig4-4-project-detail.png",
        "caption": "图4-4 项目详情与报销申请任务流图",
    },
    {
        "heading": "4.2.5 项目流程详情模块设计",
        "paragraphs": [
            "项目流程详情模块主要用于展示项目在不同流程阶段中的状态信息，面向系统管理员与教师用户开放。用户可在该页面中查看流程阶段、审批意见、审批人及提交时间，并根据当前阶段要求上传材料、提交阶段数据以及查看或下载已上传附件。",
            "由于该模块既承担流程查看功能，也承担阶段提交功能，因此在任务流中包含阶段选择、材料上传和状态反馈等多个关键节点。该模块任务流图如图4-5所示：",
        ],
        "image": "fig4-5-process-detail.png",
        "caption": "图4-5 项目流程详情模块任务流图",
    },
    {
        "heading": "4.2.6 项目流程审批模块设计",
        "paragraphs": [
            "项目流程审批模块面向系统管理员以及一级、二级、三级审批人开放，用于对项目流程中待审批的阶段任务进行处理。审批人员可根据流程阶段、审批状态、审批人及时间范围等条件定位待办任务，并对流程执行通过或驳回操作。",
            "该模块的核心流程在于任务筛选、审批意见录入与审批结果回写，能够满足项目流程多节点审批的业务需求。该模块任务流图如图4-6所示：",
        ],
        "image": "fig4-6-process-approval.png",
        "caption": "图4-6 项目流程审批模块任务流图",
    },
    {
        "heading": "4.2.7 申请审批模块设计",
        "paragraphs": [
            "申请审批模块主要面向系统管理员以及多级审批角色开放，用于处理项目相关申请业务。用户可在该模块中按照项目名称、申请人、审批状态和当前节点进行筛选，并查看申请详情、审批记录以及执行通过或拒绝等操作。",
            "为保证申请处理过程的完整性，该模块在任务流中体现了从条件检索、详情查看到审批执行的完整路径。该模块任务流图如图4-7所示：",
        ],
        "image": "fig4-7-application-approval.png",
        "caption": "图4-7 申请审批模块任务流图",
    },
    {
        "heading": "4.2.8 成果管理模块设计",
        "paragraphs": [
            "成果管理模块主要面向系统管理员与教师用户，用于对科研成果信息进行统一管理。用户可按成果名称、成果类型和级别进行查询，并可执行成果导出、成果详情查看及成果删除等操作，以实现成果数据的统一维护。",
            "成果管理的业务流程以成果检索和结果处理为主，支持在列表层面完成常见管理操作。该模块任务流图如图4-8所示：",
        ],
        "image": "fig4-8-achievement.png",
        "caption": "图4-8 成果管理模块任务流图",
    },
    {
        "heading": "4.2.9 报销管理模块设计",
        "paragraphs": [
            "报销管理模块主要面向系统管理员与财务管理员开放，用于对项目报销单据进行集中审批和管理。用户可根据项目名称、申请人、报销类型和状态进行筛选，并查看报销详情、预览附件以及执行通过或拒绝操作。",
            "为了满足报销审核流程需要，该模块在任务流中体现了列表筛选、明细查看和审批处理三类核心操作。该模块任务流图如图4-9所示：",
        ],
        "image": "fig4-9-reimbursement.png",
        "caption": "图4-9 报销管理模块任务流图",
    },
    {
        "heading": "4.2.10 用户管理模块设计",
        "paragraphs": [
            "用户管理模块仅对系统管理员开放，主要用于完成部门结构维护与用户信息管理。管理员可在部门树中进行节点维护，并在用户列表中执行查询、新增、编辑、删除和导出等操作，从而实现系统账号的集中管控。",
            "用户管理模块既包含部门维度的筛选流程，也包含用户对象的增删改查流程。该模块任务流图如图4-10所示：",
        ],
        "image": "fig4-10-user.png",
        "caption": "图4-10 用户管理模块任务流图",
    },
    {
        "heading": "4.2.11 角色管理模块设计",
        "paragraphs": [
            "角色管理模块同样仅对系统管理员开放，用于对系统中的角色名称、角色标识及角色描述进行统一维护。管理员可查看角色列表，并执行新增角色、编辑角色和删除角色等操作，为系统权限分配提供基础数据支撑。",
            "该模块任务流主要围绕角色列表和角色维护操作展开，整体路径清晰、操作集中。该模块任务流图如图4-11所示：",
        ],
        "image": "fig4-11-role.png",
        "caption": "图4-11 角色管理模块任务流图",
    },
]


def clear_document(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def add_paragraph(doc: Document, text: str, style: str, align=None, bold: bool = False) -> None:
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold


def add_picture(doc: Document, path: Path) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = 0
    p.paragraph_format.left_indent = 0
    p.paragraph_format.right_indent = 0
    p.paragraph_format.space_before = 0
    p.paragraph_format.space_after = 0
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run()
    run.add_picture(str(path), width=Cm(14.4))


def main() -> None:
    doc = Document(str(TEMPLATE_PATH))
    clear_document(doc)

    add_paragraph(doc, "第 4 章 概要设计", "Heading 1", WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "4.1 简介", "Heading 2")
    add_paragraph(
        doc,
        "在实际进行高校科研管理系统的开发过程中，需要重点关注系统功能结构、权限控制与业务操作流程的组织方式。为了使系统在项目管理、审批管理、成果管理、报销管理以及用户权限维护等方面保持清晰稳定的交互逻辑，需要在概要设计阶段对各功能模块的任务流进行统一梳理。",
        "Normal",
    )
    add_paragraph(
        doc,
        "本章结合前文需求分析内容，对系统中的主要功能模块进行任务流设计。通过任务流图可以更直观地展示用户在不同模块中的操作路径以及各功能节点之间的衔接关系，为后续系统实现与测试提供依据。",
        "Normal",
    )

    add_paragraph(doc, "4.2 功能结构设计", "Heading 2")
    add_paragraph(
        doc,
        "根据系统的实际业务需求，可将前端功能进一步划分为登录、首页工作台、项目管理、项目详情与报销申请、项目流程详情、项目流程审批、申请审批、成果管理、报销管理、用户管理与角色管理等模块。各模块既具备明确的职责边界，又在业务流程上相互衔接，形成完整的科研管理业务闭环。",
        "Normal",
    )

    for item in MODULES:
        add_paragraph(doc, item["heading"], "Heading 3")
        for paragraph in item["paragraphs"]:
            add_paragraph(doc, paragraph, "Normal")
        add_picture(doc, IMAGES_DIR / item["image"])
        add_paragraph(doc, item["caption"], "Normal", WD_ALIGN_PARAGRAPH.CENTER, bold=True)

    add_paragraph(doc, "4.3 本章小结", "Heading 2")
    add_paragraph(
        doc,
        "本章依据系统业务需求，对高校科研管理系统的主要功能模块进行了任务流设计。通过对业务路径的梳理，可以看出系统在项目、审批、成果、报销以及权限管理等核心模块中均形成了较为完整的操作闭环，为后续的详细设计与系统实现奠定了基础。",
        "Normal",
    )

    doc.save(str(OUTPUT_PATH))
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
